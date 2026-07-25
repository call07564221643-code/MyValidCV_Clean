"""Evidence-grounded CV structure and export helpers.

This module deliberately avoids inventing content. It identifies common CV
sections, replaces only the professional summary, and preserves the remaining
source text for explicit candidate review.
"""

from io import BytesIO
import re

from docx import Document
from docx.shared import Pt


SECTION_ALIASES = {
    "summary": {
        "profile", "professional profile", "personal profile", "summary",
        "professional summary", "career summary", "personal statement",
    },
    "skills": {
        "skills", "key skills", "core skills", "technical skills",
        "competencies", "core competencies",
    },
    "experience": {
        "experience", "work experience", "employment", "employment history",
        "professional experience", "career history", "work history",
    },
    "education": {
        "education", "education and training", "academic background",
        "qualifications", "academic qualifications",
    },
    "certifications": {
        "certifications", "certificates", "licences", "licenses",
        "professional development", "training",
    },
    "projects": {"projects", "key projects", "selected projects"},
    "interests": {"interests", "hobbies", "interests and activities"},
    "references": {"references"},
}

SECTION_TITLES = {
    "header": "Contact Details",
    "summary": "Professional Summary",
    "skills": "Key Skills",
    "experience": "Professional Experience",
    "education": "Education and Qualifications",
    "certifications": "Certifications and Training",
    "projects": "Selected Projects",
    "interests": "Interests",
    "references": "References",
    "other": "Additional Information",
}


def _normalise_heading(line):
    return re.sub(r"[^a-z0-9& ]+", "", line.lower()).strip()


def identify_section_heading(line):
    """Return a canonical section key when a short line is a known heading."""
    stripped = (line or "").strip()
    if not stripped or len(stripped) > 48:
        return None
    normalised = _normalise_heading(stripped.rstrip(":"))
    for key, aliases in SECTION_ALIASES.items():
        if normalised in aliases:
            return key
    return None


def parse_cv_sections(cv_text):
    """Split source CV text into ordered sections without discarding content."""
    lines = [line.rstrip() for line in re.split(r"\r?\n", cv_text or "")]
    sections = []
    current = {"key": "header", "title": SECTION_TITLES["header"], "lines": []}

    def flush():
        nonlocal current
        cleaned = [line for line in current["lines"]]
        while cleaned and not cleaned[0].strip():
            cleaned.pop(0)
        while cleaned and not cleaned[-1].strip():
            cleaned.pop()
        if cleaned:
            sections.append({**current, "lines": cleaned, "text": "\n".join(cleaned)})

    for line in lines:
        key = identify_section_heading(line)
        if key:
            flush()
            current = {"key": key, "title": SECTION_TITLES[key], "lines": []}
        else:
            current["lines"].append(line)
    flush()

    if not sections and (cv_text or "").strip():
        sections.append({
            "key": "other",
            "title": SECTION_TITLES["other"],
            "lines": [(cv_text or "").strip()],
            "text": (cv_text or "").strip(),
        })
    return sections


def find_evidence_citations(cv_text, terms, limit=6):
    """Return exact source lines supporting matched terms."""
    def term_matches_line(term, line):
        term_words = re.findall(r"[a-z0-9]+", term.lower())
        line_words = re.findall(r"[a-z0-9]+", line.lower())
        if not term_words:
            return False
        for term_word in term_words:
            if len(term_word) >= 7:
                if not any(word.startswith(term_word[:6]) for word in line_words):
                    return False
            elif term_word not in line_words:
                return False
        return True

    citations = []
    seen = set()
    for raw_line in re.split(r"\r?\n|(?<=[.!?])\s+", cv_text or ""):
        line = re.sub(r"\s+", " ", raw_line).strip(" -•*\t")
        if not 20 <= len(line) <= 260:
            continue
        matched_terms = [term for term in terms if term_matches_line(term, line)]
        key = line.lower()
        if matched_terms and key not in seen:
            citations.append({"text": line, "terms": matched_terms[:4]})
            seen.add(key)
        if len(citations) >= limit:
            break
    return citations


def build_structured_cv_draft(cv_text, target_role, proposed_summary, matched, missing):
    """Create a clean full-text draft while preserving source-CV sections."""
    sections = parse_cv_sections(cv_text)
    original_summary = ""
    summary_replaced = False
    output_sections = []

    for section in sections:
        if section["key"] == "summary" and not summary_replaced:
            original_summary = section["text"]
            output_sections.append({
                "key": "summary",
                "title": SECTION_TITLES["summary"],
                "text": proposed_summary,
            })
            summary_replaced = True
        else:
            output_sections.append({
                "key": section["key"],
                "title": section["title"],
                "text": section["text"],
            })

    if not summary_replaced:
        insert_at = 1 if output_sections and output_sections[0]["key"] == "header" else 0
        output_sections.insert(insert_at, {
            "key": "summary",
            "title": SECTION_TITLES["summary"],
            "text": proposed_summary,
        })

    full_text_parts = []
    for section in output_sections:
        if section["key"] == "header":
            full_text_parts.append(section["text"])
        else:
            full_text_parts.append(f"{section['title']}\n{section['text']}")

    citations = find_evidence_citations(cv_text, matched)
    return {
        "target_role": target_role,
        "sections": output_sections,
        "full_text": "\n\n".join(part for part in full_text_parts if part.strip()).strip(),
        "original_summary": original_summary or "No clearly labelled professional summary was detected.",
        "proposed_summary": proposed_summary,
        "citations": citations,
        "changes": [
            {
                "section": "Professional Summary",
                "reason": f"Focuses the opening profile on verified evidence relevant to {target_role}.",
                "original": original_summary or "No clearly labelled summary was detected.",
                "proposed": proposed_summary,
                "citations": citations[:3],
            }
        ],
        "verified_skills": matched[:10],
        "excluded_gaps": missing[:10],
    }


def cv_text_to_docx(content, title="Tailored CV"):
    """Render an edited plain-text CV draft as a clean DOCX document."""
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    document.core_properties.title = title

    known_headings = set(SECTION_TITLES.values())
    lines = (content or "").splitlines()
    first_content = True
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            document.add_paragraph()
            continue
        if first_content:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.bold = True
            run.font.size = Pt(16)
            first_content = False
        elif line in known_headings:
            document.add_heading(line, level=1)
        elif line.startswith(("- ", "• ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
