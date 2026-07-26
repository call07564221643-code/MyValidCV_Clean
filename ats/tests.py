from io import BytesIO
from types import SimpleNamespace

from docx import Document
from django.core.files.uploadedfile import SimpleUploadedFile
from django.contrib.auth.models import User
from django.test import SimpleTestCase, TestCase, override_settings
from django.urls import reverse

from subscriptions.models import CustomerSubscription, SubscriptionPlan

from .bullet_rewriting import apply_bullet_decisions, extract_experience_bullets, propose_safe_bullet
from .cv_drafting import (
    build_structured_cv_draft,
    cv_text_to_docx,
    is_legacy_generated_cv,
    parse_cv_sections,
)
from .forms import ATSAnalysisForm, MultipleFileField, validate_document
from .models import (
    ATSResult,
    CV,
    CVBulletSuggestion,
    EnterpriseBatch,
    EnterpriseCandidateResult,
    GeneratedCV,
    JobRole,
)
from .scoring import (
    _build_evidence_map,
    _calculate_confidence,
    _detect_mandatory_qualifications,
    _extract_requirement_terms,
    _has_requirement_evidence,
    calculate_score_details,
    validate_job_description,
)
from .views import (
    _validate_public_job_url,
    build_cover_letter,
    build_cv_draft_preview,
    build_interview_plan,
    build_reliability_guidance,
    build_truth_gate_summary,
    calculate_score,
    enterprise_daily_usage,
    humanize_requirement_term,
    format_document_heading,
    resolve_generated_cv_draft,
)


class UploadAndUrlSecurityTests(SimpleTestCase):
    def test_private_job_url_is_rejected(self):
        with self.assertRaises(ValueError):
            _validate_public_job_url("http://127.0.0.1/internal")

    def test_non_http_job_url_is_rejected(self):
        with self.assertRaises(ValueError):
            _validate_public_job_url("file:///etc/passwd")

    def test_nonstandard_job_url_port_is_rejected(self):
        with self.assertRaises(ValueError):
            _validate_public_job_url("https://example.com:8443/jobs/1")

    def test_executable_disguised_upload_extension_is_rejected(self):
        field = MultipleFileField()
        upload = SimpleUploadedFile("candidate.exe", b"not a cv")
        with self.assertRaisesMessage(Exception, "PDF, DOCX or TXT"):
            field.clean([upload])

    def test_more_than_fifteen_files_is_rejected_before_processing(self):
        field = MultipleFileField()
        uploads = [SimpleUploadedFile(f"cv-{index}.txt", b"cv") for index in range(16)]
        with self.assertRaisesMessage(Exception, "no more than 15"):
            field.clean(uploads)

    def test_pdf_extension_with_binary_content_is_rejected(self):
        upload = SimpleUploadedFile("candidate.pdf", b"MZ executable content")
        with self.assertRaisesMessage(Exception, "valid PDF signature"):
            validate_document(upload)

    def test_binary_txt_file_is_rejected(self):
        upload = SimpleUploadedFile("candidate.txt", b"text\x00binary")
        with self.assertRaisesMessage(Exception, "binary data"):
            validate_document(upload)


class EnterpriseWorkspaceTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("enterprise-owner", password="password")
        self.job_role = JobRole.objects.create(
            user=self.user,
            title="Operations Manager",
            company="Example Ltd",
            description="Operations management, reporting and stakeholder communication.",
        )
        self.batch = EnterpriseBatch.objects.create(
            user=self.user,
            job_role=self.job_role,
            title="Operations shortlist",
        )
        EnterpriseCandidateResult.objects.create(
            batch=self.batch,
            candidate_name="Alex Candidate",
            cv_file="enterprise_cvs/alex.txt",
            score=72,
            matched_skills="Operations management, reporting",
            missing_skills="Budget ownership",
            recommendation="Review the supporting evidence with the hiring panel.",
            rank=1,
        )

    def test_enterprise_report_is_concise_responsive_evidence_workspace(self):
        self.client.force_login(self.user)
        response = self.client.get(reverse("enterprise_report", args=[self.batch.id]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Enterprise evidence workspace")
        self.assertContains(response, "Candidate evidence comparison")
        self.assertContains(response, "enterprise-mobile-list")
        self.assertContains(response, "Advisory results—not automated hiring decisions")
        self.assertContains(response, "Above 50% alignment")
        self.assertNotContains(response, "Cover-letter draft")

    def test_daily_usage_reuses_existing_candidate_rows(self):
        self.assertEqual(enterprise_daily_usage(self.user), 1)


class StructuredCVDraftTests(SimpleTestCase):
    source_cv = """Alex Example
alex@example.com

Professional Summary
Operations professional supporting business teams.

Key Skills
Communication
Reporting

Work Experience
Business Operations Officer
Prepared weekly reports and communicated operational risks to stakeholders.

Education
Business Administration Diploma
"""

    def test_parser_preserves_common_cv_sections(self):
        sections = parse_cv_sections(self.source_cv)
        self.assertEqual(
            [section["key"] for section in sections],
            ["header", "summary", "skills", "experience", "education"],
        )

    def test_full_draft_replaces_summary_and_preserves_source_history(self):
        draft = build_structured_cv_draft(
            self.source_cv,
            "Operations Manager",
            "Operations professional with verified reporting and communication experience.",
            ["reporting", "communication"],
            ["budget ownership"],
        )

        self.assertIn("Operations professional with verified", draft["full_text"])
        self.assertIn("Business Operations Officer", draft["full_text"])
        self.assertIn("Business Administration Diploma", draft["full_text"])
        self.assertNotIn("Operations professional supporting business teams.", draft["full_text"])
        self.assertTrue(draft["citations"])
        self.assertIn("weekly reports", draft["citations"][0]["text"])

    def test_docx_export_produces_office_document(self):
        document = cv_text_to_docx(self.source_cv, "Alex CV")
        self.assertTrue(document.startswith(b"PK"))
        self.assertGreater(len(document), 1000)
        parsed = Document(BytesIO(document))
        self.assertEqual(parsed.paragraphs[0].text, "Alex Example")
        self.assertIn("Professional Summary", [paragraph.text for paragraph in parsed.paragraphs])

    def test_parser_understands_decorated_and_alternative_headings(self):
        sections = parse_cv_sections(
            "Alex Example\n\n1. ABOUT ME\nProfile text.\n\n"
            "02 - PROFESSIONAL BACKGROUND\nRole history.\n\n"
            "EDUCATION / QUALIFICATIONS\nDiploma.\n\nAWARDS & ACHIEVEMENTS\nAward."
        )
        self.assertEqual(
            [section["key"] for section in sections],
            ["header", "summary", "experience", "education", "achievements"],
        )

    def test_legacy_generated_format_is_detected(self):
        self.assertTrue(is_legacy_generated_cv("Targeted CV Section Draft for HR Coordinator"))
        self.assertFalse(is_legacy_generated_cv(self.source_cv))


class BulletRewritingTests(SimpleTestCase):
    def test_experience_bullets_are_extracted_with_source_evidence(self):
        cv_text = """Alex Example

Professional Experience
Operations Officer | Example Ltd | 2022 - Present
Responsible for managing weekly reporting for senior stakeholders
Maintained an Excel tracker and reduced unresolved cases by 18%.

Education
Diploma
"""
        suggestions = extract_experience_bullets(
            cv_text,
            ["reporting", "stakeholder", "excel"],
        )
        self.assertEqual(len(suggestions), 2)
        self.assertTrue(suggestions[0]["proposed_text"].startswith("Managed"))
        self.assertIn("reporting", suggestions[0]["evidence_terms"])
        self.assertFalse(suggestions[0]["has_measure"])
        self.assertTrue(suggestions[1]["has_measure"])

    def test_safe_proposal_does_not_add_new_facts(self):
        proposed, changed = propose_safe_bullet(
            "Responsible for coordinating interview schedules"
        )
        self.assertTrue(changed)
        self.assertEqual(proposed, "Coordinated interview schedules.")
        self.assertNotIn("improved", proposed.lower())

    def test_only_accepted_or_edited_decisions_are_applied(self):
        suggestions = [
            SimpleNamespace(
                status="accepted",
                original_text="Managed weekly reports.",
                proposed_text="Managed weekly operational reports.",
                edited_text="",
            ),
            SimpleNamespace(
                status="rejected",
                original_text="Supported meetings.",
                proposed_text="Coordinated meetings.",
                edited_text="",
            ),
        ]
        rebuilt, applied = apply_bullet_decisions(
            "Managed weekly reports.\nSupported meetings.",
            suggestions,
        )
        self.assertEqual(applied, 1)
        self.assertIn("Managed weekly operational reports.", rebuilt)
        self.assertIn("Supported meetings.", rebuilt)

    def test_document_headings_are_normalised_centrally(self):
        self.assertEqual(
            format_document_heading("alex morgan - hr cv"),
            "Alex Morgan - HR CV",
        )
        self.assertEqual(
            format_document_heading("hr coordinator"),
            "HR Coordinator",
        )


class EditableCVDraftEndpointTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("plus-user", password="password")
        plan = SubscriptionPlan.objects.create(
            code="plus",
            name="Plus",
            includes_generated_cv=True,
            monthly_analysis_limit=20,
        )
        CustomerSubscription.objects.create(user=self.user, plan=plan, status="active")
        self.source_cv_text = (
            "Candidate CV\n\nProfessional Summary\nOperations professional.\n\n"
            "Professional Experience\nOperations Officer | Example Ltd | 2022 - Present\n"
            "Responsible for managing weekly reporting for senior stakeholders\n"
            "Maintained an Excel tracker and reduced unresolved cases by 18%.\n\n"
            "Education\nBusiness Diploma | 2020"
        )
        self.cv = CV.objects.create(
            user=self.user,
            title="Candidate CV",
            file="cvs/candidate.txt",
            original_filename="candidate.txt",
            mime_type="text/plain",
            file_data=self.source_cv_text.encode(),
        )
        self.result = ATSResult.objects.create(
            user=self.user,
            cv=self.cv,
            job_title="Operations Manager",
            job_description="Operations Manager role requiring reporting and communication.",
            score=70,
        )
        self.generated_cv = GeneratedCV.objects.create(
            user=self.user,
            original_cv=self.cv,
            ats_result=self.result,
            title="Candidate CV tailored",
            content="Candidate CV\n\nProfessional Summary\n" + ("Verified operations experience. " * 8),
        )
        self.client.force_login(self.user)

    def test_owner_can_save_and_export_docx(self):
        revised = "Candidate CV\n\nProfessional Summary\n" + ("Verified reporting experience. " * 8)
        save_response = self.client.post(
            reverse("save_generated_cv_draft", args=[self.result.id]),
            {"cv_draft_content": revised},
        )
        self.assertRedirects(
            save_response,
            f"{reverse('ats_result', args=[self.result.id])}#full-cv-workspace",
            fetch_redirect_response=False,
        )
        self.generated_cv.refresh_from_db()
        self.assertEqual(self.generated_cv.content, revised.strip())
        self.assertEqual(
            resolve_generated_cv_draft(self.result, self.generated_cv)["draft_state"],
            "candidate_edited",
        )

        export_response = self.client.get(reverse("download_generated_cv_docx", args=[self.result.id]))
        self.assertEqual(export_response.status_code, 200)
        self.assertEqual(
            export_response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertTrue(export_response.content.startswith(b"PK"))

    def test_other_user_cannot_edit_draft(self):
        other = User.objects.create_user("other-user", password="password")
        self.client.force_login(other)
        response = self.client.post(
            reverse("save_generated_cv_draft", args=[self.result.id]),
            {"cv_draft_content": "Unauthorised change " * 20},
        )
        self.assertEqual(response.status_code, 404)

    def test_legacy_draft_is_upgraded_without_overwriting_saved_record(self):
        legacy = (
            "Targeted CV Section Draft for Operations Manager\n\n"
            "Proposed Replacement Summary\nLegacy advisory wording.\n\n"
            "Original CV Content Reference\nOld reference."
        )
        self.generated_cv.content = legacy
        self.generated_cv.save(update_fields=["content"])

        draft = resolve_generated_cv_draft(self.result, self.generated_cv)
        self.assertEqual(draft["draft_state"], "legacy_upgraded")
        self.assertNotIn("Targeted CV Section Draft", draft["editable_content"])
        self.generated_cv.refresh_from_db()
        self.assertEqual(self.generated_cv.content, legacy)

        response = self.client.get(reverse("download_generated_cv", args=[self.result.id]))
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "Targeted CV Section Draft")

    def test_owner_can_review_and_apply_bullet_decisions(self):
        start = self.client.post(reverse("start_bullet_review", args=[self.result.id]))
        self.assertRedirects(
            start,
            f"{reverse('ats_result', args=[self.result.id])}#stage-2-bullet-review",
            fetch_redirect_response=False,
        )
        suggestions = list(self.result.bullet_suggestions.all())
        self.assertEqual(len(suggestions), 2)

        first = suggestions[0]
        decision = self.client.post(
            reverse("decide_bullet_suggestion", args=[self.result.id, first.id]),
            {"decision": "accepted"},
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )
        self.assertEqual(decision.status_code, 200)
        self.assertEqual(decision.json()["status"], "accepted")
        self.assertEqual(decision.json()["summary"]["approved"], 1)
        first.refresh_from_db()
        self.assertEqual(first.status, "accepted")

        self.generated_cv.content = self.source_cv_text
        self.generated_cv.save(update_fields=["content"])
        applied = self.client.post(
            reverse("apply_bullet_review", args=[self.result.id]),
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )
        self.assertEqual(applied.status_code, 200)
        self.assertTrue(applied.json()["saved"])
        self.assertEqual(applied.json()["applied"], 1)
        self.assertIn(first.proposed_text, applied.json()["content"])
        self.generated_cv.refresh_from_db()
        self.assertIn(first.proposed_text, self.generated_cv.content)

    def test_other_user_cannot_decide_owned_bullet(self):
        suggestion = CVBulletSuggestion.objects.create(
            user=self.user,
            ats_result=self.result,
            position=0,
            fingerprint="a" * 64,
            original_text="Managed weekly reporting for senior stakeholders.",
            proposed_text="Managed weekly reporting for senior stakeholders.",
            rationale="Keeps verified wording.",
        )
        other = User.objects.create_user("bullet-other", password="password")
        self.client.force_login(other)
        response = self.client.post(
            reverse("decide_bullet_suggestion", args=[self.result.id, suggestion.id]),
            {"decision": "accepted"},
        )
        self.assertEqual(response.status_code, 404)


class CoverLetterTests(SimpleTestCase):
    def test_letter_addresses_recipient_and_keeps_drafting_note_outside_content(self):
        user = SimpleNamespace(
            username="alex",
            get_full_name=lambda: "Alex Morgan",
        )
        result = SimpleNamespace(
            job_title="Operations Manager",
            job_role=SimpleNamespace(company="Northstar Logistics"),
        )

        letter = build_cover_letter(
            user,
            result,
            ["leadership", "operations"],
            "Led daily operations and improved delivery performance by 15%.",
        )

        self.assertTrue(letter.startswith("Dear Hiring Manager,"))
        self.assertIn("Yours sincerely,\nAlex Morgan", letter)
        self.assertNotIn("Dear Mr Alex", letter)
        self.assertNotIn("Draft note:", letter)


class ATSV2Tests(TestCase):
    def test_generic_advert_grammar_is_not_promoted_to_requirements(self):
        terms = _extract_requirement_terms(
            (
                "Essential capabilities include recruitment and onboarding for a growing team. "
                "The successful candidate will coordinate the complete process."
            ),
            "HR Coordinator",
            ["recruitment", "onboarding"],
        )
        self.assertIn("recruitment", terms)
        self.assertIn("onboarding", terms)
        for generic in (
            "essential", "capabilities", "include", "growing", "successful",
            "coordinate", "complete",
        ):
            self.assertNotIn(generic, terms)

    def test_qualification_mentions_are_classified_without_false_verification(self):
        scenarios = (
            ("Currently researching CIPD Level 3 but not yet enrolled.", "not_held"),
            ("Currently studying towards CIPD Level 3.", "training"),
            ("CIPD membership lapsed in 2023.", "expired"),
            ("Awarded CIPD Level 3 in 2023.", "verified"),
        )
        for cv_text, expected_status in scenarios:
            with self.subTest(cv_text=cv_text):
                evidence = _build_evidence_map(
                    cv_text,
                    ["cipd"],
                    [],
                    ["cipd"],
                    "CIPD qualification is preferred.",
                )
                self.assertEqual(evidence[0]["status"], expected_status)

        self.assertFalse(
            _has_requirement_evidence(
                "cipd",
                "Currently researching CIPD Level 3 but not yet enrolled.",
                ["cipd"],
            )
        )
        self.assertTrue(
            _has_requirement_evidence(
                "cipd",
                "Awarded CIPD Level 3 in 2023.",
                ["cipd"],
            )
        )

    def test_short_or_placeholder_job_advert_is_rejected(self):
        valid, reason = validate_job_description("Job advert URL: https://example.com/job")
        self.assertFalse(valid)
        self.assertIn("too short", reason)

    def test_job_advert_requires_meaningful_role_signals(self):
        valid, reason = validate_job_description("This is general website text. " * 20)
        self.assertFalse(valid)
        self.assertIn("complete job advert", reason)

    def test_complete_job_advert_is_accepted(self):
        advert = (
            "Software Developer role. Responsibilities include building and testing web services. "
            "Required skills include Python, Django, SQL, Git, communication and API development. "
            "Candidates must have relevant software experience and knowledge of secure deployment."
        )
        valid, reason = validate_job_description(advert)
        self.assertTrue(valid, reason)

    def test_keyword_only_match_is_not_treated_as_verified_evidence(self):
        details = calculate_score_details(
            "Profile\nSkills: Python, Django, SQL.\nExperience\nGeneral office support.\n"
            "Education\nDiploma\nalex@example.com\n" + ("Additional profile text. " * 25),
            "Developer role. Required skills include Python, Django and SQL. "
            "Responsibilities include building, testing and deploying secure web APIs. " * 3,
            "Django Developer",
        )
        django_evidence = next(item for item in details["evidence_map"] if item["term"] == "django")
        self.assertEqual(django_evidence["status"], "mentioned")
        self.assertEqual(django_evidence["strength"], "keyword only")
        self.assertEqual(details["model_version"], "2.1")
        self.assertIn("evidence", details["score_components"])
        self.assertIn("format", details["score_components"])

    def test_negated_or_preferred_qualification_is_not_mandatory(self):
        qualification = SimpleNamespace(
            normalized_name="forklift licence",
            is_license=True,
            terms=lambda: ["forklift licence"],
        )
        self.assertEqual(
            _detect_mandatory_qualifications(
                "A forklift licence is preferred but not required.",
                [qualification],
            ),
            [],
        )

    def test_candidate_can_record_truth_gate_confirmation(self):
        user = User.objects.create_user("candidate", "candidate@example.com", "password")
        cv = CV.objects.create(user=user, title="Candidate CV", file="cvs/candidate.txt")
        result = ATSResult.objects.create(
            user=user,
            cv=cv,
            job_title="Developer",
            job_description=(
                "Developer role with responsibilities for web services. Required skills include "
                "Python, Django, SQL and testing. Candidates must have software experience. " * 2
            ),
            metrics={
                "model_version": "2.1",
                "score_components": {},
                "evidence_map": [{"term": "django", "status": "mentioned"}],
                "requirement_groups": {},
                "confidence": {},
            },
        )
        self.client.force_login(user)

        response = self.client.post(
            reverse("ats_result", args=[result.id]),
            {
                "requirement": "django",
                "evidence_action": "confirmed",
                "truth_gate_item": "1",
            },
        )

        expected_url = f"{reverse('ats_result', args=[result.id])}#truth-gate-item-1"
        self.assertRedirects(response, expected_url)
        result.refresh_from_db()
        self.assertEqual(result.metrics["candidate_confirmations"]["django"], "confirmed")
        rendered = self.client.get(reverse("ats_result", args=[result.id]))
        self.assertEqual(rendered.status_code, 200)
        self.assertContains(rendered, "Evidence decision studio")
        self.assertContains(rendered, 'id="truth-gate"')
        self.assertContains(rendered, 'class="truth-gate-scroll"')
        self.assertContains(rendered, "selected-confirmed")
        self.assertContains(rendered, "I have this experience")
        self.assertContains(rendered, "Interview Studio: Developer")
        self.assertContains(rendered, "Your next action")
        self.assertContains(rendered, "Rate your experience")
        self.assertContains(rendered, "mvcv-feedback-visible")
        self.assertContains(rendered, "private thank-you code")
        self.assertNotContains(rendered, '<select name="public_identity">')
        self.assertContains(rendered, "--truth-card-bg:")
        self.assertContains(rendered, "--truth-selected-bg:")
        self.assertContains(rendered, ".truth-status.verified { background: #1f3374; color: #fff; }")
        self.assertContains(rendered, 'aria-label="Your report journey"')
        self.assertContains(rendered, "requirement-title")
        self.assertContains(rendered, "truth-ring")
        self.assertContains(rendered, 'class="truth-gate-title"')
        self.assertContains(rendered, "font-size: 1.65rem")
        self.assertContains(rendered, "font-size: 1.45rem")
        self.assertContains(rendered, "How reliable is this check?")
        self.assertNotContains(rendered, "Assessment confidence:")
        self.assertContains(rendered, "--semantic-low: #fa3737")
        self.assertContains(rendered, "--semantic-high: #73d179")
        self.assertContains(rendered, "--semantic-neutral: #a6a4a4")
        self.assertContains(rendered, "--mvcv-brand-primary: #9fb8ff")
        self.assertContains(rendered, 'html[data-bs-theme="dark"] .letter-output')
        self.assertContains(rendered, "text-align: left")
        self.assertNotContains(rendered, "text-align: justify")
        self.assertContains(rendered, "Evidence to strengthen first")
        self.assertNotContains(rendered, "Proposed summary wording")


class TruthGateGuidanceTests(SimpleTestCase):
    def test_requirement_labels_are_human_readable(self):
        self.assertEqual(humanize_requirement_term("skillsproficiency"), "Skills proficiency")
        self.assertEqual(humanize_requirement_term("project_management"), "Project management")

    def test_cv_preview_uses_detected_role_and_filters_generic_terms(self):
        result = SimpleNamespace(
            job_title="Imported Job Role",
            metrics={"taxonomy": {"detected_role": "Operations Manager"}},
            cv=SimpleNamespace(title="Candidate CV"),
        )

        preview = build_cv_draft_preview(
            result,
            ["communication", "office", "software", "administrative", "various"],
            ["budgeting"],
            "Managed administrative operations and communicated with business stakeholders.",
        )

        self.assertEqual(preview["target_role"], "Operations Manager")
        self.assertIn("communication, administrative", preview["summary"])
        self.assertNotIn("Imported Job Role", preview["summary"])
        self.assertNotIn("various", preview["summary"])

    def test_completion_summary_reflects_candidate_actions(self):
        summary = build_truth_gate_summary([
            {"term": "python", "candidate_action": "confirmed"},
            {"term": "testing", "candidate_action": "training"},
            {"term": "aws", "candidate_action": "not_have"},
        ])

        self.assertTrue(summary["complete"])
        self.assertEqual(summary["completion"], 100)
        self.assertEqual(summary["confirmed"], 1)
        self.assertEqual(summary["training"], 1)
        self.assertEqual(summary["not_have"], 1)
        self.assertEqual(summary["confirmed_percentage"], 33)
        self.assertEqual(summary["training_percentage"], 33)
        self.assertEqual(summary["gap_percentage"], 33)
        self.assertIn("unsupported requirements", summary["next_action"])

    def test_interview_plan_uses_role_and_truth_gate_evidence(self):
        plan = build_interview_plan(
            [{"term": "safeguarding", "status": "mentioned", "candidate_action": "training"}],
            "Care Assistant",
        )

        self.assertEqual(plan["role"], "Care Assistant")
        self.assertIn("Care Assistant", plan["standard"][0]["prompt"])
        self.assertIn("currently learning about Safeguarding", plan["tailored"][0]["prompt"])

    def test_missing_role_template_cannot_receive_high_reliability_label(self):
        evidence_map = [{"term": f"skill {index}", "passage": "Used in a role."} for index in range(5)]
        score, label, reasons = _calculate_confidence(
            "A" * 800,
            "B" * 600,
            [f"skill {index}" for index in range(5)],
            evidence_map,
            "",
        )

        self.assertEqual(score, 85)
        self.assertEqual(label, "Medium")
        self.assertIn("No curated role template matched this advert.", reasons)

    def test_reliability_guidance_uses_plain_language_without_percentage(self):
        guidance = build_reliability_guidance({
            "label": "Medium",
            "reasons": ["No curated role template matched this advert."],
        })

        self.assertEqual(guidance["status"], "Review carefully")
        self.assertIn("could not identify the exact job type", guidance["message"])
        self.assertEqual(guidance["details"], ["The exact job type could not be identified."])


class ATSScoringTests(TestCase):
    def test_admin_cv_is_capped_for_dentist_role(self):
        cv_text = """
        Office Administrator
        Profile: organised administrator with five years of experience.
        Skills: administration, scheduling, records management, data entry,
        customer service, Excel, document control and compliance.
        Experience: managed appointments, supported customers, coordinated
        reports, improved filing accuracy by 20%, and trained two new staff.
        Education: business administration diploma.
        """
        dentist_job = """
        Dentist required for a busy dental clinic. Responsibilities include
        clinical assessment, oral health diagnosis, treatment planning,
        restorative dentistry, patient care, infection control, radiography,
        x-ray review and compliance with GDC standards.
        """

        score, matched, missing, recommendation = calculate_score(cv_text, dentist_job, "Dentist")

        self.assertLess(score, 50)
        self.assertIn("dentist", missing)
        self.assertIn("dental", missing)
        self.assertIn("High role mismatch", recommendation)


    def test_admin_cv_gets_partial_fit_for_airport_admin_role(self):
        cv_text = """
        Office Administrator
        Profile: organised administrator with five years of experience.
        Skills: administration, scheduling, records management, data entry,
        customer service, Excel, document control and compliance.
        Experience: managed appointments, supported customers, coordinated
        reports, improved filing accuracy by 20%, and trained two new staff.
        Education: business administration diploma.
        """
        airport_admin_job = """
        Airport administrator required to support airport operations. The role
        includes passenger service records, scheduling, document control,
        compliance, coordination with terminal teams, flight administration,
        email correspondence and accurate data entry.
        """

        score, matched, missing, recommendation = calculate_score(cv_text, airport_admin_job, "Airport Administrator")

        self.assertGreaterEqual(score, 50)
        self.assertLess(score, 90)
        self.assertIn("administration", matched)
        self.assertIn("airport", missing)
        self.assertNotIn("High role mismatch", recommendation)

    def test_software_cv_scores_well_for_software_role(self):
        cv_text = """
        Backend Software Engineer
        Skills: Python, Django, SQL, PostgreSQL, API development, Git, testing.
        Experience: developed REST APIs, improved database performance by 30%,
        deployed services, fixed defects, and worked with product teams.
        Education: computer science degree.
        """
        software_job = """
        Senior Django Developer required. Must have Python, Django, REST API,
        PostgreSQL, SQL, Git, testing experience, database optimisation and
        ability to deploy backend services.
        """

        score, matched, missing, recommendation = calculate_score(cv_text, software_job, "Senior Django Developer")

        self.assertGreaterEqual(score, 70)
        self.assertIn("django", matched)
        self.assertNotIn("High role mismatch", recommendation)

    def test_software_cv_is_capped_for_finance_role(self):
        cv_text = """
        Backend Software Engineer
        Skills: Python, Django, SQL, PostgreSQL, API development, Git, testing.
        Experience: developed REST APIs, improved database performance by 30%,
        deployed services, fixed defects, and worked with product teams.
        Education: computer science degree.
        """
        accountant_job = """
        Management Accountant required. Responsibilities include month-end
        accounts, reconciliations, budgeting, forecasting, variance analysis,
        balance sheet control, VAT returns, payroll journals and financial
        reporting using accounting software.
        """

        score, matched, missing, recommendation = calculate_score(cv_text, accountant_job, "Management Accountant")

        self.assertLess(score, 50)
        self.assertIn("accountant", missing)
        self.assertIn("High role mismatch", recommendation)


@override_settings(SECURE_SSL_REDIRECT=False)
class CVManagementTests(TestCase):
    """Prove front-end CV update/delete CRUD and object ownership rules."""

    def setUp(self):
        self.owner = User.objects.create_user("cv-owner", "owner@example.com", "password")
        self.other_user = User.objects.create_user("other-user", "other@example.com", "password")
        self.cv = CV.objects.create(
            user=self.owner,
            title="Original CV",
            file="cvs/original.txt",
            original_filename="original.txt",
            file_size=100,
        )

    def test_manage_page_reads_only_the_users_cvs(self):
        CV.objects.create(user=self.other_user, title="Private CV", file="cvs/private.txt")
        self.client.force_login(self.owner)
        response = self.client.get(reverse("upload_cv"))
        self.assertContains(response, "Original CV")
        self.assertNotContains(response, "Private CV")
        self.assertContains(response, reverse("update_cv", args=[self.cv.public_id]))
        self.assertContains(response, reverse("delete_cv", args=[self.cv.public_id]))

    def test_owner_can_update_cv_title(self):
        self.client.force_login(self.owner)
        response = self.client.post(
            reverse("update_cv", args=[self.cv.public_id]),
            {"title": "Django Developer CV"},
        )
        self.assertRedirects(response, reverse("upload_cv"))
        self.cv.refresh_from_db()
        self.assertEqual(self.cv.title, "Django Developer CV")

    def test_user_cannot_update_another_users_cv(self):
        self.client.force_login(self.other_user)
        response = self.client.post(
            reverse("update_cv", args=[self.cv.public_id]),
            {"title": "Changed without permission"},
        )
        self.assertEqual(response.status_code, 404)
        self.cv.refresh_from_db()
        self.assertEqual(self.cv.title, "Original CV")

    def test_delete_requires_confirmation_post(self):
        self.client.force_login(self.owner)
        url = reverse("delete_cv", args=[self.cv.public_id])
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Yes, delete CV")
        self.assertTrue(CV.objects.filter(pk=self.cv.pk).exists())

        response = self.client.post(url)
        self.assertRedirects(response, reverse("upload_cv"))
        self.assertFalse(CV.objects.filter(pk=self.cv.pk).exists())

    def test_deleting_cv_cascades_to_related_result(self):
        result = ATSResult.objects.create(
            user=self.owner,
            cv=self.cv,
            job_title="Developer",
            job_description="Python and Django developer",
        )
        self.client.force_login(self.owner)
        self.client.post(reverse("delete_cv", args=[self.cv.public_id]))
        self.assertFalse(ATSResult.objects.filter(pk=result.pk).exists())

    def test_user_cannot_delete_another_users_cv(self):
        self.client.force_login(self.other_user)
        response = self.client.post(reverse("delete_cv", args=[self.cv.public_id]))
        self.assertEqual(response.status_code, 404)
        self.assertTrue(CV.objects.filter(pk=self.cv.pk).exists())
