"""
ATS MVP Engine: Extract, parse, and match CV against job description.
Returns realistic scores only when both inputs are valid.
"""

import re
import string
from pathlib import Path
from typing import Dict, Tuple, Optional, List
import io

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Optional NLP & fuzzy matching libraries
try:
    import spacy
    try:
        # Try to load a small English model if available
        try:
            _nlp = spacy.load('en_core_web_sm')
            HAS_SPACY = True
        except Exception:
            # Model not installed; fall back to loading blank English pipeline
            try:
                _nlp = spacy.blank('en')
                HAS_SPACY = True
            except Exception:
                _nlp = None
                HAS_SPACY = False
    except Exception:
        _nlp = None
        HAS_SPACY = False
except ImportError:
    spacy = None
    _nlp = None
    HAS_SPACY = False

try:
    from rapidfuzz import process, fuzz
    HAS_RAPIDFUZZ = True
except Exception:
    process = None
    fuzz = None
    HAS_RAPIDFUZZ = False


class CVExtractionError(Exception):
    """Raised when CV extraction fails validation."""
    pass


class JobDescriptionError(Exception):
    """Raised when job description validation fails."""
    pass


class ATSEngine:
    """
    Extract and match CV against job description.
    Produces realistic scores based on actual content analysis.
    """

    # Skill keywords - common technical and soft skills
    SKILL_KEYWORDS = {
        # Programming languages
        'python', 'javascript', 'java', 'csharp', 'c#', 'php', 'ruby', 'go', 'rust',
        'typescript', 'kotlin', 'swift', 'objective-c', 'scala', 'perl', 'r',
        'matlab', 'sql', 'html', 'css', 'xml', 'json', 'yaml',
        
        # Frameworks and libraries
        'django', 'flask', 'fastapi', 'react', 'vue', 'angular', 'spring',
        'spring boot', 'express', 'nodejs', 'node.js', 'asp.net', 'dot net',
        'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy',
        
        # Databases
        'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'oracle',
        'sql server', 'cassandra', 'dynamodb', 'firebase',
        
        # Cloud and DevOps
        'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'jenkins',
        'gitlab', 'github', 'ci/cd', 'terraform', 'ansible', 'prometheus',
        
        # Tools and platforms
        'git', 'linux', 'windows', 'macos', 'ios', 'android', 'salesforce',
        'jira', 'confluence', 'slack', 'figma', 'photoshop', 'illustrator',
        
        # Soft skills
        'communication', 'leadership', 'teamwork', 'problem solving', 'time management',
        'project management', 'agile', 'scrum', 'kanban', 'cross-functional',
        'collaboration', 'documentation', 'testing', 'debugging',
        
        # Data and analytics
        'data analysis', 'machine learning', 'deep learning', 'nlp', 'computer vision',
        'statistics', 'data visualization', 'tableau', 'power bi', 'analytics',
    }

    # Experience signals (years, certifications, etc.)
    EXPERIENCE_PATTERNS = {
        'years': r'(\d+)\+?\s*(?:years?|yrs?|years?\s+of|experience)(?:\s+(?:in|with))?',
        'certification': r'(?:certified?|certification|cert|badge|credential)',
        'education': r'(?:degree|diploma|bachelor|master|phd|mba|b\.s\.|m\.s\.|b\.a\.|m\.a\.)',
        'achievement': r'(?:led|managed|built|developed|designed|architected|implemented)',
    }

    # Qualification keywords
    QUALIFICATION_KEYWORDS = {
        'certification', 'certified', 'credential', 'degree', 'diploma', 'bachelor',
        'master', 'phd', 'associate', 'licensed', 'cpa', 'pmp', 'aws', 'azure',
        'google', 'scrum', 'agile', 'itil', 'cissp', 'security+',
    }

    def __init__(self):
        """Initialize the ATS engine."""
        pass

    # ============================================================================
    # FILE EXTRACTION
    # ============================================================================

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        if not HAS_PYPDF:
            raise CVExtractionError("pypdf library not installed. Install with: pip install pypdf")

        try:
            text = ""
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise CVExtractionError(f"Failed to extract text from PDF: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        if not HAS_DOCX:
            raise CVExtractionError("python-docx library not installed. Install with: pip install python-docx")

        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            raise CVExtractionError(f"Failed to extract text from DOCX: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            raise CVExtractionError(f"Failed to read TXT file: {str(e)}")

    def extract_text_from_upload(self, uploaded_file) -> str:
        """Extract text from Django uploaded file object (no disk storage)."""
        file_name = uploaded_file.name.lower()

        try:
            if file_name.endswith('.pdf'):
                if not HAS_PYPDF:
                    raise CVExtractionError("pypdf library not installed")
                reader = pypdf.PdfReader(uploaded_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text

            elif file_name.endswith('.docx'):
                if not HAS_DOCX:
                    raise CVExtractionError("python-docx library not installed")
                doc = Document(uploaded_file)
                return "\n".join([p.text for p in doc.paragraphs])

            elif file_name.endswith('.txt'):
                try:
                    return uploaded_file.read().decode('utf-8')
                except UnicodeDecodeError:
                    return uploaded_file.read().decode('latin-1')

            else:
                raise CVExtractionError(f"Unsupported file format: {file_name}")

        except CVExtractionError:
            raise
        except Exception as e:
            raise CVExtractionError(f"Failed to extract text: {str(e)}")

    # ============================================================================
    # VALIDATION
    # ============================================================================

    def validate_cv_text(self, text: str) -> Tuple[bool, str]:
        """
        Validate if text looks like a real CV.
        Returns (is_valid, reason).
        """
        text_lower = text.lower()
        text_lines = text.strip().split('\n')

        # Check minimum length (CV should have substantial content)
        if len(text) < 500:
            return False, "CV is too short (minimum 500 characters)"

        # Check for common CV sections
        cv_indicators = [
            'experience', 'education', 'skill', 'summary', 'objective',
            'qualification', 'employment', 'work', 'project'
        ]
        has_cv_section = any(indicator in text_lower for indicator in cv_indicators)

        if not has_cv_section:
            return False, "No CV sections found (missing: experience, education, skills)"

        # Check for email or contact info
        has_contact = bool(re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text))
        if not has_contact:
            return False, "No contact information (email) found"

        # Check line variety (should have headers, content, etc.)
        if len(text_lines) < 10:
            return False, "CV structure appears incomplete"

        return True, "Valid CV"

    def validate_job_description_text(self, text: str) -> Tuple[bool, str]:
        """
        Validate if text looks like a real job description.
        Returns (is_valid, reason).
        """
        text_lower = text.lower()

        # Check minimum length
        if len(text) < 300:
            return False, "Job description is too short (minimum 300 characters)"

        # Check for common job description indicators
        jd_indicators = [
            'requirement', 'responsibility', 'skill', 'experience', 'education',
            'qualification', 'about', 'role', 'position', 'job', 'duty',
            'task', 'preferred', 'desired', 'must have', 'should have'
        ]
        has_jd_section = any(indicator in text_lower for indicator in jd_indicators)

        if not has_jd_section:
            return False, "No job description sections found (missing: requirements, responsibilities, skills)"

        # Check for job title or role name
        text_lines = text.strip().split('\n')
        if len(text_lines) < 5:
            return False, "Job description structure appears incomplete"

        return True, "Valid Job Description"

    # ============================================================================
    # EXTRACTION LOGIC
    # ============================================================================

    def extract_skills(self, text: str) -> List[str]:
        """Extract skills mentioned in text."""
        text_lower = text.lower()
        found_skills = set()

        # 1) Direct keyword matching (exact)
        for skill in self.SKILL_KEYWORDS:
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(skill.title())

        # 2) spaCy noun-chunk and entity extraction for multi-word skills (if available)
        if HAS_SPACY and _nlp is not None:
            try:
                doc = _nlp(text)
                for chunk in doc.noun_chunks:
                    phrase = chunk.text.strip()
                    if len(phrase) > 2 and len(phrase.split()) <= 4:
                        found_skills.add(phrase.title())
                # Also include named entities that look like technologies
                for ent in doc.ents:
                    ent_text = ent.text.strip()
                    if len(ent_text) > 2 and len(ent_text.split()) <= 3:
                        found_skills.add(ent_text.title())
            except Exception:
                # If spaCy parsing fails, ignore and continue
                pass

        # 3) Fuzzy match remaining candidate words against known skill keywords (if available)
        if HAS_RAPIDFUZZ:
            try:
                # Prepare a candidate list of words and short phrases
                candidates = set(re.findall(r"\b[A-Za-z0-9+#.+-]{3,}\b(?:\s+[A-Za-z0-9+#.+-]{3,})?", text))
                # Use RapidFuzz to extract best matches
                for cand in candidates:
                    matches = process.extract(str(cand), list(self.SKILL_KEYWORDS), scorer=fuzz.token_sort_ratio, limit=3)
                    for match in matches:
                        label, score, _ = match
                        if score >= 80:
                            found_skills.add(label.title())
            except Exception:
                pass

        # Normalize and return
        normalized = set()
        for s in found_skills:
            s_clean = re.sub(r'\s+', ' ', s).strip()
            if s_clean:
                normalized.add(s_clean)

        return sorted(normalized)

    def extract_experience_signals(self, text: str) -> Dict:
        """Extract experience-related signals from text."""
        text_lower = text.lower()
        signals = {
            'years_of_experience': 0,
            'has_certifications': False,
            'has_education': False,
            'achievements_count': 0,
        }

        # Extract years of experience
        years_match = re.search(self.EXPERIENCE_PATTERNS['years'], text_lower)
        if years_match:
            signals['years_of_experience'] = int(years_match.group(1))

        # Check for certifications
        if re.search(self.EXPERIENCE_PATTERNS['certification'], text_lower):
            signals['has_certifications'] = True

        # Check for education
        if re.search(self.EXPERIENCE_PATTERNS['education'], text_lower):
            signals['has_education'] = True

        # Count achievement verbs
        achievement_matches = re.findall(self.EXPERIENCE_PATTERNS['achievement'], text_lower)
        signals['achievements_count'] = len(achievement_matches)

        return signals

    def extract_qualifications(self, text: str) -> List[str]:
        """Extract qualifications mentioned in text."""
        text_lower = text.lower()
        found_qualifications = []

        for qual in self.QUALIFICATION_KEYWORDS:
            pattern = r'\b' + re.escape(qual) + r'\b'
            if re.search(pattern, text_lower):
                found_qualifications.append(qual.title())

        return sorted(list(set(found_qualifications)))

    def extract_keywords(self, text: str, top_n: int = 20) -> List[str]:
        """Extract important keywords from text."""
        # Remove common words
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'is', 'are', 'was', 'were', 'be',
            'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'should',
            'could', 'may', 'might', 'must', 'can', 'that', 'this', 'these',
            'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'as', 'if',
            'about', 'up', 'into', 'through', 'during', 'before', 'after',
            'above', 'below', 'between', 'under', 'again', 'further', 'then',
            'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all',
            'each', 'every', 'both', 'few', 'more', 'most', 'other', 'some',
            'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than',
            'too', 'very', 'just', 'don', 's', 't', 'now', 'd', 'll', 'm', 've',
            'what', 'which', 'who', 'whom', 'because', 'any', 'our', 'their',
            'me', 'him', 'her', 'us', 'them', 'your', 'his', 'mine', 'yours',
        }

        # Extract words
        words = re.findall(r'\b[a-z]+\b', text.lower())
        
        # Filter and count
        word_freq = {}
        for word in words:
            if word not in stop_words and len(word) > 3:
                word_freq[word] = word_freq.get(word, 0) + 1

        # Sort by frequency
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, freq in sorted_words[:top_n]]

    # ============================================================================
    # MATCHING & SCORING
    # ============================================================================

    def match_skills(self, cv_skills: List[str], jd_skills: List[str]) -> Tuple[List[str], List[str]]:
        """Match CV skills against job description skills."""
        cv_skills_lower = [s.lower() for s in cv_skills]
        jd_skills_lower = [s.lower() for s in jd_skills]

        matched = [s for s, sl in zip(cv_skills, cv_skills_lower) if sl in jd_skills_lower]
        missing = [s for s in jd_skills if s.lower() not in cv_skills_lower]

        return matched, missing

    def match_keywords(self, cv_keywords: List[str], jd_keywords: List[str]) -> Tuple[List[str], List[str]]:
        """Match CV keywords against job description keywords."""
        cv_keywords_lower = [k.lower() for k in cv_keywords]
        jd_keywords_lower = [k.lower() for k in jd_keywords]

        matched = [k for k, kl in zip(cv_keywords, cv_keywords_lower) if kl in jd_keywords_lower]
        missing = [k for k in jd_keywords if k.lower() not in cv_keywords_lower]

        return matched, missing

    def calculate_scores(
        self,
        cv_skills: List[str],
        jd_skills: List[str],
        cv_keywords: List[str],
        jd_keywords: List[str],
        cv_experience: Dict,
        jd_experience: Dict,
        matched_skills: List[str],
        matched_keywords: List[str],
    ) -> Dict:
        """Calculate realistic ATS matching scores."""
        scores = {}

        # Skills score (0-100)
        if jd_skills:
            skills_score = min(100, int((len(matched_skills) / len(jd_skills)) * 100))
        else:
            skills_score = 50
        scores['skills_score'] = skills_score

        # Keywords score (0-100)
        if jd_keywords:
            keywords_score = min(100, int((len(matched_keywords) / len(jd_keywords)) * 100))
        else:
            keywords_score = 50
        scores['keywords_score'] = keywords_score

        # Experience score (0-100)
        cv_years = cv_experience.get('years_of_experience', 0)
        experience_score = min(100, (cv_years * 10) + (cv_experience.get('achievements_count', 0) * 5))
        scores['experience_score'] = experience_score

        # Qualification score (0-100)
        qualification_points = 0
        if cv_experience.get('has_education'):
            qualification_points += 40
        if cv_experience.get('has_certifications'):
            qualification_points += 30
        scores['qualification_score'] = min(100, qualification_points)

        # Format score (0-100) - based on structure
        format_score = 75  # Good base for valid CVs
        scores['format_score'] = format_score

        # ATS score (average of key metrics)
        ats_score = int((skills_score + keywords_score + scores['qualification_score'] + format_score) / 4)
        scores['ats_score'] = ats_score

        # Overall score (weighted average)
        overall_score = int(
            (skills_score * 0.35) +
            (keywords_score * 0.25) +
            (experience_score * 0.20) +
                (scores['qualification_score'] * 0.15) +
            (format_score * 0.05)
        )
        scores['overall_score'] = overall_score

        return scores

    def get_confidence_level(self, cv_text_len: int, jd_text_len: int) -> Tuple[int, int]:
        """Calculate confidence levels for CV and job description."""
        # CV confidence: based on text length and completeness
        # More content = higher confidence
        cv_confidence = min(100, int((cv_text_len / 5000) * 100))
        if cv_confidence < 40:
            cv_confidence = 40

        # JD confidence: based on text length
        jd_confidence = min(100, int((jd_text_len / 2000) * 100))
        if jd_confidence < 50:
            jd_confidence = 50

        return cv_confidence, jd_confidence

    # ============================================================================
    # MAIN PUBLIC METHOD
    # ============================================================================

    def analyse(self, cv_text: str, jd_text: str) -> Dict:
        """
        Main analysis method: extract, validate, parse, and match.
        Returns comprehensive analysis report or error information.
        """
        result = {
            'success': False,
            'error': None,
            'data': {}
        }

        # Step 1: Validate inputs
        cv_valid, cv_reason = self.validate_cv_text(cv_text)
        if not cv_valid:
            result['error'] = f"Invalid CV: {cv_reason}"
            return result

        jd_valid, jd_reason = self.validate_job_description_text(jd_text)
        if not jd_valid:
            result['error'] = f"Invalid Job Description: {jd_reason}"
            return result

        # Step 2: Extract information
        cv_skills = self.extract_skills(cv_text)
        cv_experience = self.extract_experience_signals(cv_text)
        cv_qualifications = self.extract_qualifications(cv_text)
        cv_keywords = self.extract_keywords(cv_text)

        jd_skills = self.extract_skills(jd_text)
        jd_experience = self.extract_experience_signals(jd_text)
        jd_qualifications = self.extract_qualifications(jd_text)
        jd_keywords = self.extract_keywords(jd_text)

        # Step 3: Match
        matched_skills, missing_skills = self.match_skills(cv_skills, jd_skills)
        matched_keywords, missing_keywords = self.match_keywords(cv_keywords, jd_keywords)

        # Step 4: Calculate scores
        scores = self.calculate_scores(
            cv_skills, jd_skills, cv_keywords, jd_keywords,
            cv_experience, jd_experience,
            matched_skills, matched_keywords
        )

        # Step 5: Calculate confidence
        cv_confidence, jd_confidence = self.get_confidence_level(len(cv_text), len(jd_text))

        # Step 6: Build result
        result['success'] = True
        result['data'] = {
            'overall_score': scores['overall_score'],
            'ats_score': scores['ats_score'],
            'skills_score': scores['skills_score'],
            'keywords_score': scores['keywords_score'],
            'experience_score': scores['experience_score'],
            'qualification_score': scores['qualification_score'],
            'format_score': scores['format_score'],
            'matched_skills': matched_skills,
            'missing_skills': missing_skills,
            'matched_keywords': matched_keywords,
            'missing_keywords': missing_keywords,
            'cv_confidence': cv_confidence,
            'job_confidence': jd_confidence,
            'cv_qualifications': cv_qualifications,
            'cv_experience_years': cv_experience.get('years_of_experience', 0),
            'cv_has_education': cv_experience.get('has_education', False),
            'cv_has_certifications': cv_experience.get('has_certifications', False),
        }

        return result


# Create a singleton instance
ats_engine = ATSEngine()
